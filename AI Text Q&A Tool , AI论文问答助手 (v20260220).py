import time
t1 = time.time()

#本软件来自：https://gemini.google.com/app/0c6575a05e2c28d0
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
import docx
from docx.oxml.ns import qn  # 用于设置中文字体
import PyPDF2
import threading
import os
from zhipuai import ZhipuAI

# 配置文件路径
CONFIG_FILE = "config.txt"

class AIQuestionSystem:
    def __init__(self, root):
        self.root = root
        # 软件标题
        self.root.title("AI文本智能问答系统 | 基于AI大模型(智谱GLM-4 + 谷歌Gemini 3) | 20260220版")
        self.root.geometry("950x980")

        self.remaining_time = 0
        self.timer_active = False  # 新增：用于判断倒计时是否已经在运行

        # 第一行：致谢与联系方式
        thanks_label = tk.Label(root,
                                text="谢天谢地，感谢AI软件工程师一龙哥哥，本软件永久免费使用，如有问题请查看帮助文档或联系(+86)18620997719(手机/微信)",
                                fg="blue", font=("微软雅黑", 9))
        thanks_label.pack(pady=5)

        # --- 1. 文章内容区域 ---
        tk.Label(root, text="1. 上传论文 / 粘贴文字", font=("微软雅黑", 10, "bold")).pack(pady=2)
        self.article_area = scrolledtext.ScrolledText(root, height=10, width=110)
        self.article_area.pack(pady=2, padx=20)

        btn_frame = tk.Frame(root)
        btn_frame.pack()
        tk.Button(btn_frame, text="上传本地文件(PDF/Word/TXT)", command=self.upload_file).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="清空输入框内容", command=lambda: self.article_area.delete("1.0", tk.END)).pack(
            side=tk.LEFT)

        # --- 2. 参数调节 ---
        param_frame = tk.LabelFrame(root, text="AI文本大模型微调参数", font=("微软雅黑", 10, "bold"), fg="black")
        param_frame.pack(pady=5, padx=20, fill="x")

        tk.Label(param_frame, text="Temperature(思维发散温度):").grid(row=0, column=0, padx=5, pady=5)
        self.temp_slider = tk.Scale(param_frame, from_=0.0, to=1.0, resolution=0.01, orient=tk.HORIZONTAL, length=220)
        self.temp_slider.set(0.50)
        self.temp_slider.grid(row=0, column=1, padx=5)

        tk.Label(param_frame, text="Top_p(只选前p%):").grid(row=0, column=2, padx=5, pady=5)
        self.top_p_slider = tk.Scale(param_frame, from_=0.0, to=1.0, resolution=0.01, orient=tk.HORIZONTAL, length=220)
        self.top_p_slider.set(0.50)
        self.top_p_slider.grid(row=0, column=3, padx=5)

        # --- 3. 提问区域 (2行) ---
        tk.Label(root, text="2. 对原文内容提问", font=("微软雅黑", 10, "bold")).pack(pady=2)
        self.question_input = scrolledtext.ScrolledText(root, height=2, width=110, font=("微软雅黑", 10))
        self.question_input.pack(pady=2, padx=20)

        # --- 4. 底部区域 (API输入 + 按钮组) ---
        bottom_action_frame = tk.Frame(root)
        bottom_action_frame.pack(pady=10, padx=20, fill="x")

        tk.Label(bottom_action_frame, text="智谱 API Key:").pack(side=tk.LEFT)
        self.api_entry = tk.Entry(bottom_action_frame, width=30, show="*")
        self.api_entry.pack(side=tk.LEFT, padx=5)
        self.load_api_key()

        self.status_label = tk.Label(bottom_action_frame, text="就绪", font=("微软雅黑", 9), fg="green", width=15)
        self.status_label.pack(side=tk.LEFT, padx=5)

        # 绿色开始按钮 (移除禁用逻辑，始终可以点击)
        self.submit_btn = tk.Button(bottom_action_frame, text="AI生成答案", command=self.handle_submit,
                                    bg="#28a745", fg="white", font=("微软雅黑", 9, "bold"), width=12)
        self.submit_btn.pack(side=tk.LEFT, padx=5)

        # 淡粉色导出按钮
        self.export_btn = tk.Button(bottom_action_frame, text="导出Word", command=self.export_to_word,
                                    bg="#ffc0cb", fg="black", font=("微软雅黑", 9, "bold"), width=12)
        self.export_btn.pack(side=tk.LEFT, padx=5)

        # --- 5. 分析结果 ---
        tk.Label(root, text="3. AI回答内容", font=("微软雅黑", 10, "bold")).pack(pady=2)
        self.result_area = scrolledtext.ScrolledText(root, height=30, width=110, bg="#fdfdfd", wrap=tk.WORD,
                                                     font=("微软雅黑", 10))
        self.result_area.pack(pady=2, padx=20)

    # --- 核心功能函数 ---

    def export_to_word(self):
        content = self.result_area.get("1.0", tk.END).strip()
        if not content or content == "正在请求 GLM-4...":
            messagebox.showwarning("提示", "当前没有可导出的分析结果！")
            return

        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word 文档", "*.docx")],
                                                 initialfile="AI分析结果报告.docx")
        if file_path:
            try:
                doc = docx.Document()
                style = doc.styles['Normal']
                style.font.name = '宋体'
                style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                heading = doc.add_heading('AI 文本分析报告', 0)
                for run in heading.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                para = doc.add_paragraph(content)
                for run in para.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

                doc.save(file_path)
                messagebox.showinfo("成功", f"文件已成功保存（宋体）至:\n{file_path}")
            except Exception as e:
                messagebox.showerror("导出失败", f"错误原因: {e}")

    def load_api_key(self):
        if os.path.exists(CONFIG_FILE):
            with open(CONFIG_FILE, "r") as f:
                self.api_entry.insert(0, f.read().strip())

    def save_api_key(self, key):
        with open(CONFIG_FILE, "w") as f:
            f.write(key)

    def upload_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("支持格式", "*.txt *.docx *.pdf")])
        if not file_path: return
        try:
            content = ""
            if file_path.endswith('.txt'):
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            elif file_path.endswith('.docx'):
                doc = docx.Document(file_path)
                content = "\n".join([p.text for p in doc.paragraphs])
            elif file_path.endswith('.pdf'):
                reader = PyPDF2.PdfReader(file_path)
                for page in reader.pages: content += (page.extract_text() or "") + "\n"
            self.article_area.delete("1.0", tk.END)
            self.article_area.insert(tk.END, content.strip())
            self.status_label.config(text="文件解析成功", fg="green")
        except Exception as e:
            messagebox.showerror("错误", f"文件读取失败: {e}")

    def handle_submit(self):
        """改进：移除倒计时拦截逻辑，允许直接点击"""
        key = self.api_entry.get().strip()
        if not key:
            messagebox.showwarning("提示", "请先输入有效的 API Key");
            return
        self.save_api_key(key)

        # 即使在倒计时，也可以直接启动新线程
        threading.Thread(target=self.analyze_with_glm, daemon=True).start()

    def start_or_reset_timer(self):
        """新增：启动或重新设置倒计时"""
        self.remaining_time = 30
        if not self.timer_active:
            self.timer_active = True
            self.update_timer()

    def update_timer(self):
        """改进：秒级更新，支持重置"""
        if self.remaining_time > 0:
            self.status_label.config(text=f"API 冷却：{self.remaining_time}s", fg="#D32F2F")
            self.remaining_time -= 1
            self.root.after(1000, self.update_timer)
        else:
            self.status_label.config(text="API 已就绪", fg="green")
            self.timer_active = False  # 计时结束，标记设为 False

    def analyze_with_glm(self):
        api_key = self.api_entry.get().strip()
        article = self.article_area.get("1.0", tk.END).strip()
        question = self.question_input.get("1.0", tk.END).strip()

        if not article or not question:
            messagebox.showwarning("提示", "内容或问题缺失");
            return

        self.status_label.config(text="AI正在努力思考...", fg="orange")
        self.result_area.delete("1.0", tk.END)
        self.result_area.insert(tk.END, "正在连接 GLM-4 接口...")

        try:
            client = ZhipuAI(api_key=api_key)
            response = client.chat.completions.create(
                model="glm-4",
                messages=[
                    {"role": "system", "content": "你是一个极其严谨的学术助手。请严格基于原文回答，严禁瞎编。"},
                    {"role": "user", "content": f"【原文】：\n{article}\n\n【问题】：\n{question}"}
                ],
                top_p=self.top_p_slider.get(),
                temperature=self.temp_slider.get()
            )
            self.result_area.delete("1.0", tk.END)
            self.result_area.insert(tk.END, response.choices[0].message.content)

            # 改进：成功返回后，无论之前状态如何，重新开始30秒倒计时
            self.start_or_reset_timer()

        except Exception as e:
            self.result_area.delete("1.0", tk.END)
            self.result_area.insert(tk.END, f"\n[调用出错] \n错误详情：{e}")
            self.status_label.config(text="调用失败", fg="red")

if __name__ == "__main__":
    root = tk.Tk()
    app = AIQuestionSystem(root)
    root.mainloop()

#完成计时
t2 = time.time()
print('总共耗时为', round((t2-t1)/60,2), '分钟。')

